"""
Generates Oz — Shark Tank submission document as a .docx file.
Run: python3 scripts/build_docx.py
Output: /root/orchestration/Oz_SharkTank_Document.docx
"""

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = '/root/orchestration/Oz_SharkTank_Document.docx'

# ── Colour palette ────────────────────────────────────────────────────────────
INDIGO   = RGBColor(0x4F, 0x46, 0xE5)
INDIGO_L = RGBColor(0xE0, 0xE7, 0xFF)
SLATE    = RGBColor(0x1E, 0x29, 0x3B)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GRAY     = RGBColor(0x6B, 0x72, 0x80)
GREEN    = RGBColor(0x05, 0x96, 0x69)
ORANGE   = RGBColor(0xEA, 0x58, 0x0C)
RED      = RGBColor(0xDC, 0x26, 0x26)
AMBER    = RGBColor(0xD9, 0x77, 0x06)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top='single', bottom='single', left='single', right='single', color='CCCCCC', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def para_color(para, color: RGBColor):
    for run in para.runs:
        run.font.color.rgb = color


def heading(doc, text, level, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def body(doc, text, bold=False, color=None, size=10, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return p


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_image_from_fig(doc, fig, width=6.5):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width))
    plt.close(fig)
    return p


def section_header(doc, text, emoji=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{emoji}  {text}' if emoji else text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = INDIGO
    return p


# ── Diagrams ──────────────────────────────────────────────────────────────────

def make_architecture_diagram():
    fig, ax = plt.subplots(figsize=(12, 8))
    fig.patch.set_facecolor('#0F172A')
    ax.set_facecolor('#0F172A')
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')

    def box(x, y, w, h, label, sublabel='', bg='#1E293B', fg='white', radius=0.3, fontsize=9):
        fancy = mpatches.FancyBboxPatch((x, y), w, h,
            boxstyle=f'round,pad=0.05,rounding_size={radius}',
            facecolor=bg, edgecolor='#334155', linewidth=1.5)
        ax.add_patch(fancy)
        cy = y + h / 2 + (0.12 if sublabel else 0)
        ax.text(x + w / 2, cy, label, ha='center', va='center',
                color=fg, fontsize=fontsize, fontweight='bold', wrap=True)
        if sublabel:
            ax.text(x + w / 2, y + h / 2 - 0.22, sublabel, ha='center', va='center',
                    color='#94A3B8', fontsize=7)

    def arrow(x1, y1, x2, y2, color='#4F46E5'):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle='->', color=color, lw=2))

    def band(y, h, label, color):
        rect = mpatches.FancyBboxPatch((0.2, y), 11.6, h,
            boxstyle='round,pad=0.05,rounding_size=0.2',
            facecolor=color, edgecolor='none', alpha=0.15)
        ax.add_patch(rect)
        ax.text(0.4, y + h - 0.18, label, color=color, fontsize=7.5,
                fontweight='bold', alpha=0.9)

    # Bands
    band(6.7, 1.1, 'USER INTERFACES', '#818CF8')
    band(4.8, 1.6, 'LEON AI ASSISTANT', '#34D399')
    band(2.8, 1.7, 'OZ CONTROL PLANE  (FastAPI · PostgreSQL · Redis)', '#F59E0B')
    band(0.3, 2.2, 'AGENT EXECUTION LAYER', '#F87171')

    # User interfaces row
    for i, (lbl, sub) in enumerate([('Leon Chat', 'WebSocket'), ('Web Dashboard', 'Tailwind UI'),
                                     ('REST API', 'FastAPI'), ('CLI', 'python oz')]):
        box(0.5 + i * 2.8, 6.85, 2.4, 0.72, lbl, sub, bg='#312E81', fg='#C7D2FE', fontsize=8)

    # Leon row
    for i, (lbl, sub) in enumerate([('NLU Pipeline', 'Intent classifier'), ('Skill Router', 'oz_skill'),
                                     ('Action LLM', 'Llama / Mistral')]):
        box(0.6 + i * 3.7, 5.05, 3.2, 0.95, lbl, sub, bg='#064E3B', fg='#6EE7B7', fontsize=8)

    # Control plane row
    for i, (lbl, sub) in enumerate([('Agent Launcher', '+ Status Poller'), ('Server Registry', 'SSH creds'),
                                     ('Secrets Vault', 'AES-256'), ('Cron Scheduler', 'Celery Beat'),
                                     ('Audit Log', 'Full trail')]):
        box(0.35 + i * 2.3, 3.0, 2.1, 0.85, lbl, sub, bg='#451A03', fg='#FCD34D', fontsize=7.5)

    # Agent layer
    for i, (lbl, sub, bg, fg) in enumerate([
        ('CF Docker\nAgent', 'list/stop/logs/exec', '#1E1B4B', '#A5B4FC'),
        ('CF Server\nHealth Agent', 'CPU/mem/disk', '#1E1B4B', '#A5B4FC'),
        ('CF Proxmox\nAgent', 'VMs/LXC/nodes', '#1E1B4B', '#A5B4FC'),
        ('oz-local\nAgent', 'SSH + bash', '#1C1917', '#D6D3D1'),
        ('+ 97 Edge\nAgents', 'planned', '#1A1A2E', '#6B7280'),
    ]):
        box(0.35 + i * 2.3, 0.45, 2.1, 1.05, lbl, sub, bg=bg, fg=fg, fontsize=7.5)

    # Arrows
    arrow(6, 6.85, 6, 6.0)   # UI → Leon
    arrow(6, 4.95, 6, 3.87)  # Leon → Control plane
    arrow(3.0, 2.95, 2.4, 1.52)  # CP → CF agents
    arrow(6.0, 2.95, 6.5, 1.52)  # CP → oz-local

    ax.text(6, 7.95, 'Oz — System Architecture', ha='center', va='center',
            color='white', fontsize=13, fontweight='bold')
    return fig


def make_request_flow_diagram():
    fig, ax = plt.subplots(figsize=(12, 5.5))
    fig.patch.set_facecolor('#0F172A')
    ax.set_facecolor('#0F172A')
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 5.5)
    ax.axis('off')

    steps = [
        ('1', 'User', '"List Docker\ncontainers on\nAITLP-371"', '#312E81', '#C7D2FE'),
        ('2', 'Leon NLU', 'Intent: docker_logs\nEntity: server=AITLP-371\nEntity: action=list', '#064E3B', '#6EE7B7'),
        ('3', 'resolveDockerTarget()', 'Matches endpoint\n"AITLP-371"\nRoutes → CF Agent', '#451A03', '#FCD34D'),
        ('4', 'CF Proxmox\n/ Docker Agent', 'Llama 4 Scout\ncalls tools\nreal API calls', '#1E1B4B', '#A5B4FC'),
        ('5', 'Leon Formatter', 'HTML card\nwith status dots\nand tables', '#1A2744', '#93C5FD'),
        ('6', 'User sees\nstructured result', '< 5 seconds\nbeautiful card\nin chat UI', '#064E3B', '#6EE7B7'),
    ]

    for i, (num, title, desc, bg, fg) in enumerate(steps):
        x = 0.4 + i * 2.1
        # Box
        rect = mpatches.FancyBboxPatch((x, 1.2), 1.85, 2.8,
            boxstyle='round,pad=0.05,rounding_size=0.2',
            facecolor=bg, edgecolor='#334155', linewidth=1.5)
        ax.add_patch(rect)
        # Number badge
        circle = plt.Circle((x + 0.925, 3.65), 0.32, color=fg, zorder=5)
        ax.add_patch(circle)
        ax.text(x + 0.925, 3.65, num, ha='center', va='center',
                color='#0F172A', fontsize=10, fontweight='bold', zorder=6)
        ax.text(x + 0.925, 2.95, title, ha='center', va='center',
                color=fg, fontsize=8, fontweight='bold')
        ax.text(x + 0.925, 1.9, desc, ha='center', va='center',
                color='#94A3B8', fontsize=7, linespacing=1.5)

        if i < len(steps) - 1:
            ax.annotate('', xy=(x + 1.95, 2.6), xytext=(x + 1.85, 2.6),
                arrowprops=dict(arrowstyle='->', color='#4F46E5', lw=2))

    ax.text(6.5, 5.15, 'Request Flow: Natural Language → Infrastructure Action',
            ha='center', va='center', color='white', fontsize=12, fontweight='bold')
    ax.text(6.5, 0.55, 'Every step is automated. The user only types plain English.',
            ha='center', va='center', color='#64748B', fontsize=9)
    return fig


def make_agent_ecosystem_diagram():
    fig, ax = plt.subplots(figsize=(12, 7))
    fig.patch.set_facecolor('#0F172A')
    ax.set_facecolor('#0F172A')
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis('off')

    ax.text(6, 6.6, 'Oz Agent Ecosystem — 101 Agents', ha='center', va='center',
            color='white', fontsize=13, fontweight='bold')
    ax.text(6, 6.15, '4 Built  ·  97 Planned  ·  All on Cloudflare Edge (sub-100ms)',
            ha='center', va='center', color='#64748B', fontsize=9)

    # Centre hub
    hub = plt.Circle((6, 3.3), 0.85, color='#4F46E5', zorder=5)
    ax.add_patch(hub)
    ax.text(6, 3.45, 'Oz', ha='center', va='center', color='white', fontsize=14, fontweight='bold', zorder=6)
    ax.text(6, 3.1, 'Control\nPlane', ha='center', va='center', color='#C7D2FE', fontsize=7, zorder=6)

    built = [
        ('🐳 CF Docker\nAgent', 2.0, 5.5, '#1E1B4B', '#A5B4FC', True),
        ('🖥️ CF Server\nHealth Agent', 9.8, 5.5, '#1E1B4B', '#A5B4FC', True),
        ('⚡ CF Proxmox\nAgent', 2.0, 1.0, '#1E1B4B', '#A5B4FC', True),
        ('🔧 oz-local\nAgent', 9.8, 1.0, '#1C1917', '#D6D3D1', True),
    ]

    categories = [
        ('🗄️ Database\n& Storage\n10 agents', 6.0, 5.8, '#1C1917', '#86EFAC'),
        ('🔒 Security\n& Compliance\n10 agents', 1.0, 3.3, '#1A1A2E', '#F9A8D4'),
        ('🌐 Web & APIs\n8 agents', 11.0, 3.3, '#1A1A2E', '#67E8F9'),
        ('📦 App Deploy\n12 agents', 3.8, 0.2, '#1A1A2E', '#FDE68A'),
        ('📊 Monitoring\n10 agents', 8.2, 0.2, '#1A1A2E', '#D9F99D'),
        ('🔀 CI/CD & Git\n7 agents', 6.0, 0.2, '#1A1A2E', '#C4B5FD'),
    ]

    def spoke_box(x, y, label, bg, fg, is_built=False):
        w, h = 1.6, 0.95
        bx, by = x - w/2, y - h/2
        ec = '#22C55E' if is_built else '#334155'
        lw = 2.5 if is_built else 1
        rect = mpatches.FancyBboxPatch((bx, by), w, h,
            boxstyle='round,pad=0.05,rounding_size=0.15',
            facecolor=bg, edgecolor=ec, linewidth=lw)
        ax.add_patch(rect)
        if is_built:
            ax.text(bx + w - 0.12, by + h - 0.12, '✅', fontsize=6, ha='right', va='top')
        ax.text(x, y, label, ha='center', va='center',
                color=fg, fontsize=7, fontweight='bold', linespacing=1.4)
        # Line to hub
        dx, dy = 6 - x, 3.3 - y
        dist = (dx**2 + dy**2) ** 0.5
        nx, ny = dx/dist, dy/dist
        ax.plot([x + nx*0.8, 6 - nx*0.85], [y + ny*0.47, 3.3 - ny*0.85],
                color='#334155', lw=1, linestyle='--', zorder=1)

    for label, x, y, bg, fg, built_flag in built:
        spoke_box(x, y, label, bg, fg, built_flag)

    for label, x, y, bg, fg in categories:
        spoke_box(x, y, label, bg, fg, False)

    # Legend
    ax.add_patch(mpatches.FancyBboxPatch((0.2, -0.05), 1.5, 0.45,
        boxstyle='round,pad=0.05,rounding_size=0.1', facecolor='#1E1B4B', edgecolor='#22C55E', linewidth=2))
    ax.text(0.95, 0.18, '✅  Built', ha='center', va='center', color='#6EE7B7', fontsize=8)
    ax.add_patch(mpatches.FancyBboxPatch((1.9, -0.05), 1.5, 0.45,
        boxstyle='round,pad=0.05,rounding_size=0.1', facecolor='#1A1A2E', edgecolor='#334155', linewidth=1))
    ax.text(2.65, 0.18, '📋  Planned', ha='center', va='center', color='#94A3B8', fontsize=8)

    return fig


def make_business_model_diagram():
    fig, ax = plt.subplots(figsize=(11, 4))
    fig.patch.set_facecolor('#0F172A')
    ax.set_facecolor('#0F172A')
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 4)
    ax.axis('off')

    tiers = [
        ('FREE\nStarter', '₹0 / mo', '3 servers\n50 runs/month\nCommunity support', '#1C1917', '#D6D3D1'),
        ('PRO', '$49 / mo', '20 servers\n500 runs/month\nAll edge agents', '#1E1B4B', '#A5B4FC'),
        ('TEAM', '$199 / mo', 'Unlimited servers\n5,000 runs/month\nSSO + Audit log', '#064E3B', '#6EE7B7'),
        ('ENTERPRISE', 'Custom', 'Self-hosted option\nCustom agents\nDedicated SLA', '#451A03', '#FCD34D'),
    ]

    for i, (name, price, desc, bg, fg) in enumerate(tiers):
        x = 0.5 + i * 2.6
        h = 3.4 if i == 2 else 3.0  # highlight Team
        by = (4 - h) / 2
        rect = mpatches.FancyBboxPatch((x, by), 2.2, h,
            boxstyle='round,pad=0.05,rounding_size=0.2',
            facecolor=bg, edgecolor=fg if i == 2 else '#334155',
            linewidth=2.5 if i == 2 else 1.5)
        ax.add_patch(rect)
        ax.text(x + 1.1, by + h - 0.4, name, ha='center', va='center',
                color=fg, fontsize=9, fontweight='bold')
        ax.text(x + 1.1, by + h - 0.85, price, ha='center', va='center',
                color='white', fontsize=11, fontweight='bold')
        ax.text(x + 1.1, by + 0.85, desc, ha='center', va='center',
                color='#94A3B8', fontsize=7.5, linespacing=1.6)
        if i == 2:
            ax.text(x + 1.1, by + h + 0.05, '★ RECOMMENDED', ha='center', va='bottom',
                    color='#6EE7B7', fontsize=7, fontweight='bold')

    ax.text(5.5, 3.85, 'SaaS Pricing Tiers', ha='center', va='center',
            color='white', fontsize=12, fontweight='bold')
    return fig


# ── Document builder ──────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run('OZ')
    run.font.size = Pt(52)
    run.font.bold = True
    run.font.color.rgb = INDIGO

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tagline.add_run('AI Agent Orchestration Platform')
    r.font.size = Pt(18)
    r.font.color.rgb = GRAY

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('"Control your entire infrastructure with plain English"')
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ('Company', 'AIT GLOBAL INC'),
        ('Contact', 'rohit.darekar@aitglobalinc.com'),
        ('Version', '1.0 — June 2026'),
        ('Category', 'AI · DevOps · Infrastructure Automation'),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = INDIGO
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = SLATE
        set_cell_bg(row.cells[0], 'F0F0FF')
        set_cell_bg(row.cells[1], 'FAFAFA')
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────────────
    section_header(doc, 'Executive Summary', '📋')
    body(doc, (
        'Oz is an AI-powered infrastructure management platform that lets any engineer '
        'manage servers, containers, deployments, and cloud resources using plain English — '
        'no memorising commands, no SSH fatigue, no tool-switching.'
    ), size=11, space_after=6)

    body(doc, 'Example queries users ask Oz today:', bold=True, size=10, space_after=2)
    for q in [
        '"How many Docker containers are running on production?"',
        '"Deploy the latest build on the web server and restart the service."',
        '"Show me the last 50 error logs from the API container."',
        '"Check the health of AITLP-371."',
        '"List all VMs on the Proxmox cluster and show me which are offline."',
    ]:
        bullet(doc, q)

    doc.add_paragraph()
    body(doc, (
        'What makes Oz different: It combines a natural-language chat interface (Leon AI), '
        'a backend orchestration engine, Cloudflare Edge AI agents (zero cold-start, globally '
        'distributed), and an encrypted secrets vault into a single deployable platform that any '
        'DevOps team can self-host today — or consume as a SaaS.'
    ), size=10, space_after=8)

    divider(doc)

    # ── 2. THE PROBLEM ────────────────────────────────────────────────────────
    section_header(doc, 'The Problem', '🚨')

    problems = [
        ('The Knowledge Gap', '#4F46E5',
         'Every infrastructure operation requires specialised CLI knowledge. New engineers spend '
         'weeks learning SSH flags and Docker commands. Senior engineers waste hours on repetitive, '
         'low-value operational tasks.'),
        ('The Context-Switching Tax', '#DC2626',
         'A typical DevOps workflow requires switching between: terminal, Docker CLI, Grafana, '
         'log aggregators, CI/CD dashboards, DNS panels, and ticketing systems. Each context '
         'switch costs 20–25 minutes of productive focus.'),
        ('The Automation Ceiling', '#D97706',
         'Existing tools (Ansible, Terraform, Bash) are powerful but brittle — they require '
         'upfront authoring, break when infrastructure changes, and cannot handle free-form '
         'operational queries. AI coding assistants help with code but cannot execute live '
         'infrastructure operations.'),
    ]
    for title, color_hex, desc in problems:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r1 = p.add_run(f'▸  {title}:  ')
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(color_hex.lstrip('#'))
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('The result: Engineering teams spend 30–40% of their time on operational toil that should be automated.')
    r.bold = True
    r.font.color.rgb = RED
    r.font.size = Pt(11)

    divider(doc)

    # ── 3. OUR SOLUTION ───────────────────────────────────────────────────────
    section_header(doc, 'Our Solution', '✅')
    body(doc, (
        'Oz solves all three problems with a single, cohesive platform. Instead of building '
        'yet another fixed automation script or another AI chatbot that only answers questions, '
        'Oz connects a natural-language understanding layer (Leon AI) to a fleet of specialised '
        'AI agents that have direct, authenticated access to your infrastructure.'
    ), size=10, space_after=6)

    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [('Problem', '4F46E5'), ('Oz Solution', '059669')]
    for i, (h, c) in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(t.rows[0].cells[i], c)
    rows = [
        ('Knowledge gap', 'Natural-language interface — ask in plain English, Oz handles the syntax'),
        ('Context switching', 'One chat for all infra — containers, servers, databases, DNS, CI/CD'),
        ('Automation ceiling', 'AI agents that reason, adapt, and execute multi-step tasks live against real infra'),
    ]
    for i, (prob, sol) in enumerate(rows):
        r = t.rows[i + 1]
        r.cells[0].text = prob
        r.cells[0].paragraphs[0].runs[0].font.bold = True
        r.cells[1].text = sol
        set_cell_bg(r.cells[0], 'F0F0FF')
    doc.add_paragraph()

    divider(doc)

    # ── 4. ARCHITECTURE DIAGRAM ───────────────────────────────────────────────
    section_header(doc, 'System Architecture', '🏗️')
    body(doc, 'The diagram below shows how the four layers of Oz work together:', size=10, space_after=4)
    add_image_from_fig(doc, make_architecture_diagram(), width=6.2)
    body(doc, (
        'Layer 1 — User Interfaces: Leon Chat (WebSocket), Web Dashboard, REST API, CLI. '
        'Layer 2 — Leon AI: NLU pipeline, skill router, action LLM. '
        'Layer 3 — Oz Control Plane: agent launcher, server registry, secrets vault, cron scheduler, audit log. '
        'Layer 4 — Agent Execution: Cloudflare Edge Workers (sub-100ms) and oz-local Docker sandboxes (for SSH tasks).'
    ), size=9, color=GRAY, space_after=6)

    divider(doc)

    # ── 5. REQUEST FLOW ───────────────────────────────────────────────────────
    section_header(doc, 'Request Flow — How It Works', '🔄')
    add_image_from_fig(doc, make_request_flow_diagram(), width=6.2)
    body(doc, (
        'Step 1: User types plain English in Leon chat. '
        'Step 2: Leon NLU pipeline identifies the intent and extracts entities (server name, container name, etc.). '
        'Step 3: resolveDockerTarget() / resolveProxmoxTarget() matches the server name to a registered endpoint. '
        'Step 4: Cloudflare Edge Worker receives the task, runs Llama 4 Scout 17B, calls real API tools. '
        'Step 5: Leon formats the output as a structured HTML card. '
        'Step 6: User sees a clean, readable result in under 5 seconds.'
    ), size=9, color=GRAY, space_after=6)

    divider(doc)

    # ── 6. KEY FEATURES ───────────────────────────────────────────────────────
    section_header(doc, 'Key Features', '⚡')
    features = [
        ('Natural Language Interface', 'No commands to memorise. Ask in plain English — Oz handles the syntax, the SSH, and the execution.'),
        ('Multi-Agent Support', 'OpenCode (free, default), Claude Code, Codex, Gemini CLI, CommandCode, oz-local, and Cloudflare edge agents. Agent-agnostic architecture.'),
        ('Cloudflare Edge AI Agents', 'Purpose-built Workers running Llama 4 Scout 17B at Cloudflare\'s global edge. Zero cold start. Sub-100ms from anywhere in the world.'),
        ('Encrypted Secrets Vault', 'SSH keys, passwords, API tokens, Proxmox tokens — all stored AES-256 encrypted. Decrypted only at agent launch time, injected as env vars, never logged.'),
        ('Server & Endpoint Registry', 'Register SSH servers and Docker/Proxmox API endpoints. Oz resolves them by name from your natural language queries.'),
        ('Cron Scheduling', 'Schedule any agent on a cron schedule. Run health checks every morning, SSL audits every week, DB backups nightly — all automatic.'),
        ('Real-time Log Streaming', 'Every agent run streams logs via WebSocket. Watch deployments unfold live in the chat UI.'),
        ('Full Audit Trail', 'Every action logged: who triggered it, which agent ran, what prompt was used, when it started and finished.'),
        ('Skills Library', 'Save reusable agent configurations as named Skills. Share best-practice setups across your team.'),
        ('Structured HTML Responses', 'Output is formatted into clean HTML cards — tables with status dots, dark log blocks, key-value metric cards — not raw text dumps.'),
    ]
    for title, desc in features:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f'✦  {title}:  ')
        r1.bold = True
        r1.font.color.rgb = INDIGO
        r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    doc.add_page_break()

    # ── 7. AGENTS BUILT ───────────────────────────────────────────────────────
    section_header(doc, 'Agents Built Today (4 of 101)', '🤖')
    add_image_from_fig(doc, make_agent_ecosystem_diagram(), width=6.2)

    built_agents = [
        ('🐳', 'cloudflare_docker_agent', 'Docker Container Management',
         'https://cf-docker-agent.rohit-darekar.workers.dev',
         '10 tools: list_containers, inspect_container, get_container_logs, container_stats, '
         'run_container, stop_container, remove_container, list_images, exec_in_container, docker_system_info',
         'List containers, fetch logs, get CPU/memory stats, exec commands, start/stop/remove containers.'),
        ('🖥️', 'cloudflare_server_health_agent', 'Server Health Monitoring',
         'https://cf-server-health-agent.rohit-darekar.workers.dev',
         '5 tools: system_info, disk_usage, list_containers, get_resource_usage, exec_in_container',
         'Full health report: OS, CPU, RAM, disk, per-container resource usage, load average, HEALTHY/WARNING/CRITICAL status.'),
        ('⚡', 'cloudflare_proxmox_agent', 'Proxmox VE Hypervisor Management',
         'https://cf-proxmox-agent.rohit-darekar.workers.dev',
         '13 tools: get_cluster_status, list_nodes, get_node_status, list_vms, get_vm_status, '
         'vm_action (start/stop/shutdown/reboot/reset), list_containers, get_container_status, '
         'container_action, list_storage, get_node_tasks, list_vm_snapshots, get_vm_config',
         'Manage VMs and LXC containers on Proxmox clusters. Check node health, storage, and task history. Start/stop VMs by name.'),
        ('🔧', 'oz-local', 'SSH + Shell Agent (Local)',
         'Docker sandbox on Oz host',
         'Runs any AI agent (OpenCode, Claude Code, Codex, Gemini CLI) inside an isolated Docker container',
         'Handles any task requiring SSH access to servers: deployments, bash commands, app logs, server health on SSH servers.'),
    ]

    for emoji, agent_id, name, url, tools, desc in built_agents:
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = t.rows[0].cells[0]
        set_cell_bg(cell, 'F8F9FF')
        set_cell_borders(cell, color='4F46E5', sz='6')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(f'{emoji}  {name}  ')
        r.bold = True
        r.font.color.rgb = INDIGO
        r.font.size = Pt(11)
        r2 = p.add_run(f'[{agent_id}]')
        r2.font.size = Pt(8)
        r2.font.color.rgb = GRAY

        p2 = cell.add_paragraph()
        p2.add_run('Endpoint: ').bold = True
        p2.add_run(url).font.size = Pt(8)

        p3 = cell.add_paragraph()
        r3a = p3.add_run('Tools: ')
        r3a.bold = True
        r3a.font.size = Pt(9)
        r3b = p3.add_run(tools)
        r3b.font.size = Pt(8)
        r3b.font.color.rgb = GRAY

        p4 = cell.add_paragraph()
        r4 = p4.add_run(desc)
        r4.font.size = Pt(9)
        r4.italic = True
        p4.paragraph_format.space_after = Pt(4)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    divider(doc)

    # ── 8. PLANNED AGENTS ────────────────────────────────────────────────────
    section_header(doc, 'Planned Agent Roadmap — 101 Agents Total', '🗺️')
    body(doc, (
        'Each new Cloudflare edge agent takes 2–4 hours to build using our standardised template. '
        'The architecture is fully templated — CF Worker + backend runner + Leon action + locale. '
        'Below is the full 10-category roadmap.'
    ), size=10, space_after=6)

    categories = [
        ('Infrastructure & Ops', 15, 2, '🖥️', ['Server health ✅', 'Process management', 'Disk analysis', 'Memory monitoring', 'Cron management', 'Systemd services', 'Log search', 'File operations', 'Package updates']),
        ('Docker & Containers', 10, 1, '🐳', ['Container management ✅', 'Docker Compose', 'Image management', 'Volume cleanup', 'Network inspection', 'Registry push/pull', 'Docker build', 'Security audit']),
        ('Proxmox VE', 1, 1, '⚡', ['VM & LXC management ✅ (13 tools: cluster, nodes, VMs, containers, storage, snapshots)']),
        ('Application & Deployment', 12, 0, '🚀', ['Deploy pipeline', 'Rollback', 'n8n workflows', 'App logs', 'nginx management', 'PM2 management', 'Feature flags', 'Smoke tests']),
        ('Database & Storage', 10, 0, '🗄️', ['PostgreSQL', 'MySQL', 'Redis', 'MongoDB', 'S3 operations', 'DB backup/restore', 'Migrations', 'SQLite']),
        ('Security & Compliance', 10, 0, '🔒', ['SSL audit', 'Firewall check', 'CVE scanning', 'SSH hardening', 'fail2ban', 'Login audit', 'Permissions audit']),
        ('Networking & DNS', 8, 0, '🌐', ['DNS lookup', 'Ping / latency', 'Port scanning', 'Traceroute', 'Bandwidth test', 'VPN status', 'WHOIS']),
        ('Monitoring & Alerting', 10, 0, '📊', ['Grafana queries', 'Prometheus PromQL', 'Sentry errors', 'Uptime monitoring', 'Anomaly detection', 'Cost estimation', 'Incident management']),
        ('AI & Productivity', 10, 0, '🧠', ['RAG over docs', 'Code review', 'Summarisation', 'Translation', 'Changelog generation', 'README generation']),
        ('Web & APIs', 8, 0, '🔗', ['Website checks', 'API testing', 'Screenshots', 'PageSpeed', 'Broken links', 'HTTP headers']),
        ('CI/CD & Git', 7, 0, '🔀', ['GitHub/GitLab status', 'PR review', 'Dependency audit', 'Release tagging', 'CI pipeline status']),
    ]

    t = doc.add_table(rows=len(categories) + 1, cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(['Category', 'Agents', 'Built', 'Key Agents']):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(t.rows[0].cells[i], '1E293B')

    for i, (cat, total, built_n, em, agents) in enumerate(categories):
        row = t.rows[i + 1]
        row.cells[0].text = f'{em}  {cat}'
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = str(total)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = '✅ ' * built_n if built_n else '📋'
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].text = ', '.join(agents[:4]) + (f'  +{len(agents)-4} more' if len(agents) > 4 else '')
        row.cells[3].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[3].paragraphs[0].runs[0].font.color.rgb = GRAY
        if built_n:
            set_cell_bg(row.cells[0], 'F0FFF4')
            set_cell_bg(row.cells[2], 'F0FFF4')

    doc.add_page_break()

    # ── 9. MARKET OPPORTUNITY ─────────────────────────────────────────────────
    section_header(doc, 'Market Opportunity', '📈')
    t = doc.add_table(rows=4, cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(['Market Segment', 'Size (2025)', 'Source']):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(t.rows[0].cells[i], '1E293B')
    market_rows = [
        ('DevOps & Platform Engineering Tools', '$8.7B', 'MarketsandMarkets'),
        ('AIOps & IT Operations AI', '$11.9B', 'IDC'),
        ('Cloud Infrastructure Management', '$29.4B', 'Gartner'),
    ]
    for i, (seg, sz, src) in enumerate(market_rows):
        t.rows[i+1].cells[0].text = seg
        t.rows[i+1].cells[1].text = sz
        t.rows[i+1].cells[1].paragraphs[0].runs[0].font.bold = True
        t.rows[i+1].cells[1].paragraphs[0].runs[0].font.color.rgb = GREEN
        t.rows[i+1].cells[2].text = src

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Combined Total Addressable Market: ~$50 Billion')
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = INDIGO

    divider(doc)

    # ── 10. BUSINESS MODEL ────────────────────────────────────────────────────
    section_header(doc, 'Business Model', '💰')
    add_image_from_fig(doc, make_business_model_diagram(), width=6.0)

    body(doc, 'Revenue Streams:', bold=True, size=10, space_after=2)
    for rev in [
        'SaaS subscriptions — Starter (free), Pro ($49/mo), Team ($199/mo), Enterprise (custom)',
        'Agent run consumption — each run consumes compute; charged above free tier limits',
        'Custom agent development — enterprise clients pay for bespoke agents for their stack',
        'Managed hosting — we run and maintain Oz for teams that prefer not to self-host',
        'Agent marketplace — community-built premium agents sold as add-ons',
    ]:
        bullet(doc, rev)

    doc.add_paragraph()
    body(doc, 'Unit Economics:', bold=True, size=10, space_after=2)
    econ = [
        ('Cloudflare Workers AI cost per edge agent run', '~$0.001–0.003'),
        ('Gross margin at Pro tier', '~75–80%'),
        ('Estimated LTV:CAC ratio', '~3:1'),
    ]
    for k, v in econ:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f'  {k}:  ')
        r1.font.size = Pt(10)
        r2 = p.add_run(v)
        r2.bold = True
        r2.font.color.rgb = GREEN
        r2.font.size = Pt(10)

    divider(doc)

    # ── 11. COMPETITIVE LANDSCAPE ─────────────────────────────────────────────
    section_header(doc, 'Competitive Landscape', '⚔️')
    comps = [
        ('Ansible / Terraform', 'Infrastructure-as-code automation', 'Requires upfront authoring and breaks on infra changes. Oz works on existing infra via natural language.'),
        ('GitHub Copilot', 'AI code assistant in the IDE', 'Suggests code but cannot execute live infrastructure operations. Oz takes real actions.'),
        ('Datadog / New Relic', 'Observability dashboards', 'Shows metrics but requires manual analysis. Oz acts — ask a question, get an action taken.'),
        ('RunBook / Airplane', 'Internal tooling platforms', 'Require workflow authoring. Oz is AI-first with no pre-built workflows needed.'),
        ('AWS Systems Manager', 'Cloud-native server management', 'AWS-only, no AI reasoning, no natural language. Oz works on any server anywhere, open-source.'),
    ]
    t = doc.add_table(rows=len(comps)+1, cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(['Competitor', 'What They Do', 'Oz Advantage']):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(t.rows[0].cells[i], '1E293B')
    for i, (comp, what, adv) in enumerate(comps):
        t.rows[i+1].cells[0].text = comp
        t.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        t.rows[i+1].cells[1].text = what
        t.rows[i+1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        t.rows[i+1].cells[2].text = adv
        t.rows[i+1].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
        t.rows[i+1].cells[2].paragraphs[0].runs[0].font.color.rgb = GREEN

    divider(doc)

    # ── 12. TECH STACK ────────────────────────────────────────────────────────
    section_header(doc, 'Tech Stack', '🛠️')
    stack = [
        ('Backend API', 'FastAPI (Python 3.12)', 'Async, auto OpenAPI docs, JWT auth'),
        ('Database', 'PostgreSQL 16', 'Relational store for users, agents, servers, secrets'),
        ('Cache / Queue', 'Redis 7 + Celery', 'Async agent execution, cron scheduling'),
        ('AI Assistant', 'Leon (open-source)', 'Production NLU pipeline, WebSocket chat, skill system'),
        ('NLU / LLM', 'Llama 3.1 / Mistral via OpenRouter', 'Intent classification and entity extraction'),
        ('Edge AI', 'Cloudflare Workers AI — Llama 4 Scout 17B', 'Sub-100ms AI inference at global edge, free tier available'),
        ('Edge Runtime', 'Cloudflare Workers (TypeScript)', 'Serverless, zero idle cost, 300+ global PoPs'),
        ('Container Runtime', 'Docker + Docker Compose', 'Reproducible deployment, agent sandboxing'),
        ('Reverse Proxy', 'nginx', 'TLS termination, Docker/Proxmox API proxying'),
        ('Security', 'JWT + bcrypt + Fernet AES', 'Stateless auth, encrypted secrets vault'),
    ]
    t = doc.add_table(rows=len(stack)+1, cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(['Layer', 'Technology', 'Why']):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(t.rows[0].cells[i], '1E293B')
    for i, (layer, tech, why) in enumerate(stack):
        t.rows[i+1].cells[0].text = layer
        t.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        t.rows[i+1].cells[1].text = tech
        t.rows[i+1].cells[1].paragraphs[0].runs[0].font.color.rgb = INDIGO
        t.rows[i+1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        t.rows[i+1].cells[2].text = why
        t.rows[i+1].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
        t.rows[i+1].cells[2].paragraphs[0].runs[0].font.color.rgb = GRAY

    doc.add_page_break()

    # ── 13. THE ASK ───────────────────────────────────────────────────────────
    section_header(doc, 'The Ask', '🤝')
    body(doc, 'We are raising a seed round to accelerate the 101-agent roadmap and go-to-market.', size=11, space_after=6)

    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    for i, h in enumerate(['Use of Funds', 'Allocation']):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        t.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(t.rows[0].cells[i], '1E293B')
    ask_rows = [
        ('Engineering (2 senior engineers)', '45%'),
        ('Complete 101-agent roadmap', '20%'),
        ('Go-to-market & sales', '25%'),
        ('Infrastructure & ops', '10%'),
    ]
    for i, (use, alloc) in enumerate(ask_rows):
        t.rows[i+1].cells[0].text = use
        t.rows[i+1].cells[1].text = alloc
        t.rows[i+1].cells[1].paragraphs[0].runs[0].font.bold = True
        t.rows[i+1].cells[1].paragraphs[0].runs[0].font.color.rgb = INDIGO
        t.rows[i+1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    section_header(doc, '12-Month Milestones', '📅')
    milestones = [
        ('M1–M3', 'Complete 25 more Cloudflare edge agents, launch public beta, first 50 customers'),
        ('M3–M6', 'Pro tier live, Stripe integration, 100 paying customers, $60K ARR'),
        ('M6–M9', 'Team tier + SSO, agent marketplace beta, 500 paying customers, $300K ARR'),
        ('M9–M12', 'Enterprise pilots, custom agent service, $1M ARR target'),
    ]
    t2 = doc.add_table(rows=len(milestones)+1, cols=2)
    t2.style = 'Table Grid'
    for i, h in enumerate(['Period', 'Milestone']):
        t2.rows[0].cells[i].text = h
        t2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        t2.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(t2.rows[0].cells[i], '1E293B')
    for i, (period, milestone) in enumerate(milestones):
        t2.rows[i+1].cells[0].text = period
        t2.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        t2.rows[i+1].cells[0].paragraphs[0].runs[0].font.color.rgb = INDIGO
        t2.rows[i+1].cells[1].text = milestone
        set_cell_bg(t2.rows[i+1].cells[0], 'F0F0FF')

    divider(doc)

    # ── 14. TEAM ─────────────────────────────────────────────────────────────
    section_header(doc, 'Team', '👥')
    p = doc.add_paragraph()
    r = p.add_run('Rohit Darekar  —  Founder & CTO, AIT GLOBAL INC')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = INDIGO
    body(doc, (
        'Full-stack engineer with deep expertise in AI systems, infrastructure automation, and product development. '
        'Built Oz end-to-end — architecture, FastAPI backend, React/Tailwind frontend, Cloudflare edge workers, '
        'Leon NLU integration, and all 4 production AI agents. '
        'Contact: rohit.darekar@aitglobalinc.com'
    ), size=10, space_after=6)

    divider(doc)

    # ── CLOSING ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run(
        'The DevOps AI tooling market is being claimed right now.\n'
        'The first platform to build a comprehensive, production-ready,\n'
        'self-hostable AI infrastructure agent fleet will own the category.\n'
        'Oz has the architecture, the working product, and the roadmap to be that platform.'
    )
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = SLATE

    doc.add_paragraph()
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = final.add_run('AIT GLOBAL INC  ·  rohit.darekar@aitglobalinc.com')
    rf.bold = True
    rf.font.color.rgb = INDIGO
    rf.font.size = Pt(11)

    doc.save(OUT)
    print(f'✅  Saved: {OUT}')


if __name__ == '__main__':
    build()
